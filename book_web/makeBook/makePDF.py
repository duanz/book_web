import os
import shutil
from book_web.utils.base_logger import logger as logging
import pdfkit
from book.models import Book, Chapter, ChapterImage, Image
from book_web.utils.common_data import BOOK_TYPE_DESC
from book.serializers import ImageSerializer
from django.conf import settings
from docx import Document
from book_web.utils.photo import split_photo_fit_kindle


class MakeMyPDF:
    def __init__(self, book_id, start_chapter_id=None, end_chapter_id=None):
        self.book_id = book_id
        self.start_chapter_id = start_chapter_id
        self.end_chapter_id = end_chapter_id
        self.book = Book.normal.get(id=book_id)
        self.content_type = self.book.book_type
        self.title = self.book.title
        self.filename = ""
        self.options = {
            "page-size": "Letter",
            "margin-top": "0.75in",
            "margin-right": "0.75in",
            "margin-bottom": "0.75in",
            "margin-left": "0.75in",
            "encoding": "UTF-8",
            "no-outline": None,
        }

    def run(self):
        if self.content_type == BOOK_TYPE_DESC.Comic:
            self.makeComicWord()
        elif self.content_type == BOOK_TYPE_DESC.Novel:
            self.makeBookPDF()

    def makeBookPDF(self):
        # 初始化txt
        if not os.path.exists(settings.MEDIA_ROOT):
            os.makedirs(settings.MEDIA_ROOT, 0o775)

        filename = os.path.join(settings.MEDIA_ROOT, self.title + ".pdf")
        if os.path.exists(filename):
            os.remove(filename)
        self.filename = filename

        # 设置章节
        start = (
            Chapter.normal.get(id=self.start_chapter_id)
            if self.start_chapter_id
            else None
        )
        end = (
            Chapter.normal.get(id=self.end_chapter_id) if self.end_chapter_id else None
        )

        chapters = Chapter.normal.filter(book=self.book)
        if start:
            chapters = chapters.filter(number__gte=start.number)
        if end:
            chapters = chapters.filter(number__lte=end.number)

        s = ""
        for chapter in chapters:
            s += chapter.title + "\n"
            s += chapter.content + "\n"
        pdfkit.from_string(s, filename, options=self.options)

        self.book.is_download = True
        self.book.save()

    def makeComicWord(self):
        # 临时文件夹
        comic_temp_path = os.path.join(settings.MEDIA_ROOT, self.title)

        part = 0
        part_size = 1024 * 1024 * 20
        current_size = 0
        pre_size = lambda cur: cur + 1024 * 1024 * 5
        # 设置章节
        chapters = Chapter.normal.filter(book=self.book)
        for chapter in chapters:
            if current_size == 0:
                # 初始化word
                part += 1
                doc = Document()
                doc.add_heading(chapter.title, level=1)
                logging.info("WORD part-{} 已经初始化".format(part))

            chapter_imgs = ChapterImage.normal.filter(chapter=chapter, book=self.book)
            if chapter_imgs:
                for img_idx, img in enumerate(chapter_imgs):
                    img_path = img.image.get_path("title")
                    img_size = os.path.getsize(img_path)
                    current_size += img_size

                    # 切割大图片临时文件夹
                    temp_path = os.path.join(
                        comic_temp_path, os.path.split(img_path)[-1].split(".")[0]
                    )

                    # 如果是大文件就分隔
                    after_split = split_photo_fit_kindle(img_path, temp_path)
                    for small_img in after_split:
                        doc.add_picture(small_img)

            if pre_size(current_size) >= part_size:
                # 保存word
                filename = os.path.join(
                    settings.MEDIA_ROOT, "{}__{}.docx".format(self.book.title, part)
                )
                if os.path.exists(filename):
                    os.remove(filename)
                doc.save(filename)
                current_size = 0
                logging.info("WORD part-{} 已经完成".format(part))

        # 删除临时文件
        shutil.rmtree(comic_temp_path)

        logging.info("word 完成")
        self.book.is_download = True
        self.book.save()
