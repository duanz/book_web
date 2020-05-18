import os
import shutil
from book_web.utils.base_logger import logger as logging

from book.models import Book, Chapter, ChapterImage, Image
from book_web.utils.common_data import BOOK_TYPE_DESC
from book.serializers import ImageSerializer
from django.conf import settings
from docx import Document
from book_web.utils.photo import split_photo_fit_kindle
# class MakeMyWord:
#     def __init__(self, content_id, content_type):
#         self.content_id = content_id
#         self.content_type = content_type
#         self.title = ""
#         self.filename = ""

#     def run(self):
#         if self.content_type == "comic":
#             self.makeComicWord()
#         elif self.content_type == 'book':
#             self.makeBookTxt()
#         pass

#     def makeBookTxt(self):
#         book_id = self.content_id
#         book_obj = Book.normal.filter(id=book_id).first()
#         self.title = book_obj.title

#         # 初始化txt
#         if not os.path.exists(settings.UPLOAD_SAVE_PATH):
#             os.makedirs(settings.UPLOAD_SAVE_PATH, 0o775)

#         filename = os.path.join(settings.UPLOAD_SAVE_PATH,
#                                 book_obj.title + '.txt')
#         if os.path.exists(filename):
#             os.remove(filename)
#         self.filename = filename
#         book = open(filename, 'w+')

#         # 设置章节
#         chapters = Chapter.normal.filter(book_id=book_id)
#         for chapter in chapters:
#             book.writelines(chapter.title + '\n')
#             book.writelines(chapter.content + '\n')

#         book.flush()
#         book.close()

#         book_obj.is_download = True
#         book_obj.save()

#     def makeComicWord(self):
#         comic_id = self.content_id
#         comic_obj = Book.normal.filter(id=comic_id).first()
#         self.title = comic_obj.title

#         # 临时文件夹
#         comic_temp_path = os.path.join(settings.UPLOAD_SAVE_PATH, self.title)

#         part = 0
#         part_size = 1024 * 1024 * 20
#         current_size = 0
#         pre_size = lambda cur: cur + 1024 * 1024 * 5
#         # 设置章节
#         chapters = Chapter.normal.filter(comic_id=comic_id)
#         for chapter in chapters:
#             if current_size == 0:
#                 # 初始化word
#                 part += 1
#                 doc = Document()
#                 doc.add_heading(chapter.title, level=1)
#                 logging.info("WORD part-{} 已经初始化".format(part))

#             chapter_imgs = ChapterImage.normal.filter(chapter=chapter,
#                                                       book=comic_obj)
#             if chapter_imgs:
#                 for img_idx, img in enumerate(chapter_imgs):
#                     img_path = img.image.get_path('title')
#                     img_size = os.path.getsize(img_path)
#                     current_size += img_size

#                     # 切割大图片临时文件夹
#                     temp_path = os.path.join(
#                         comic_temp_path,
#                         os.path.split(img_path)[-1].split('.')[0])

#                     # 如果是大文件就分隔
#                     after_split = split_photo_fit_kindle(img_path, temp_path)
#                     for small_img in after_split:
#                         doc.add_picture(small_img)

#             if pre_size(current_size) >= part_size:
#                 # 保存word
#                 filename = os.path.join(
#                     settings.UPLOAD_SAVE_PATH,
#                     '{}__{}.docx'.format(comic_obj.title, part))
#                 if os.path.exists(filename):
#                     os.remove(filename)
#                 doc.save(filename)
#                 current_size = 0
#                 logging.info("WORD part-{} 已经完成".format(part))

#         # 删除临时文件
#         shutil.rmtree(comic_temp_path)

#         logging.info("word 完成")
#         comic_obj.is_download = True
#         comic_obj.save()


class MakeMyWord:
    def __init__(self, book_id, start_chapter_id=None, end_chapter_id=None):
        self.book_id = book_id
        self.start_chapter_id = start_chapter_id
        self.end_chapter_id = end_chapter_id
        self.book = Book.normal.get(id=book_id)
        self.content_type = self.book.book_type
        self.title = self.book.title
        self.filename = ""

    def run(self):
        if self.content_type == BOOK_TYPE_DESC.Comic:
            self.makeComicWord()
        elif self.content_type == BOOK_TYPE_DESC.Novel:
            self.makeBookTxt()

    def makeBookTxt(self):
        # 初始化txt
        if not os.path.exists(settings.UPLOAD_SAVE_PATH):
            os.makedirs(settings.UPLOAD_SAVE_PATH, 0o775)

        filename = os.path.join(settings.UPLOAD_SAVE_PATH, self.title + '.txt')
        if os.path.exists(filename):
            os.remove(filename)
        self.filename = filename

        with open(filename, 'w+', encoding="utf-8") as book_handler:
            # 设置章节
            start = Chapter.normal.get(
                id=self.start_chapter_id) if self.start_chapter_id else None
            end = Chapter.normal.get(
                id=self.end_chapter_id) if self.end_chapter_id else None

            chapters = Chapter.normal.filter(book=self.book)
            if start:
                chapters = chapters.filter(number__gte=start.number)
            if end:
                chapters = chapters.filter(number__lte=end.number)

            for chapter in chapters:
                book_handler.writelines(chapter.title + '\n')
                book_handler.writelines(chapter.content + '\n')

        self.book.is_download = True
        self.book.save()

    def makeComicWord(self):
        # 临时文件夹
        comic_temp_path = os.path.join(settings.UPLOAD_SAVE_PATH, self.title)

        part = 0
        part_size = 1024 * 1024 * 20
        current_size = 0
        pre_size = lambda cur: cur + 1024 * 1024 * 5
        # 设置章节
        chapters = Chapter.normal.filter(book=self.book)
        for chapter in chapters:
            if current_size == 0:
                # 初始化word
                part += 1
                doc = Document()
                doc.add_heading(chapter.title, level=1)
                logging.info("WORD part-{} 已经初始化".format(part))

            chapter_imgs = ChapterImage.normal.filter(chapter=chapter,
                                                      book=self.book)
            if chapter_imgs:
                for img_idx, img in enumerate(chapter_imgs):
                    img_path = img.image.get_path('title')
                    img_size = os.path.getsize(img_path)
                    current_size += img_size

                    # 切割大图片临时文件夹
                    temp_path = os.path.join(
                        comic_temp_path,
                        os.path.split(img_path)[-1].split('.')[0])

                    # 如果是大文件就分隔
                    after_split = split_photo_fit_kindle(img_path, temp_path)
                    for small_img in after_split:
                        doc.add_picture(small_img)

            if pre_size(current_size) >= part_size:
                # 保存word
                filename = os.path.join(
                    settings.UPLOAD_SAVE_PATH,
                    '{}__{}.docx'.format(self.book.title, part))
                if os.path.exists(filename):
                    os.remove(filename)
                doc.save(filename)
                current_size = 0
                logging.info("WORD part-{} 已经完成".format(part))

        # 删除临时文件
        shutil.rmtree(comic_temp_path)

        logging.info("word 完成")
        self.book.is_download = True
        self.book.save()
